import os
import argparse
import pandas as pd
from jinja2 import Environment, FileSystemLoader
import glob
import re
from datetime import datetime
import json


def _relpath_posix(path: str, start: str) -> str:
    try:
        return os.path.relpath(path, start).replace(os.sep, "/")
    except Exception:
        return path


def _extract_tree_from_readme(readme_path: str) -> str:
    try:
        with open(readme_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [ln.rstrip("\n") for ln in f.readlines()]
        tree_lines = []
        for ln in lines:
            if any(tok in ln for tok in ("|--", "├──", "└──", "│   ", "──")):
                tree_lines.append(ln)
        return "\n".join(tree_lines).strip()
    except Exception:
        return ""


def _build_deliverables_tree_from_fs(input_dir: str, max_depth: int = 4) -> str:
    input_dir = os.path.abspath(input_dir)
    root_depth = input_dir.rstrip(os.sep).count(os.sep)
    out = [os.path.basename(input_dir.rstrip(os.sep)) + "/"]

    for dirpath, dirnames, filenames in os.walk(input_dir):
        depth = dirpath.rstrip(os.sep).count(os.sep) - root_depth
        if depth >= max_depth:
            dirnames[:] = []
            continue

        rel = os.path.relpath(dirpath, input_dir)
        if rel == ".":
            indent = ""
        else:
            indent = "  " * depth
            out.append(f"{indent}{os.path.basename(dirpath)}/")

        indent_files = "  " * (depth + (0 if rel == "." else 1))
        for fn in sorted(filenames)[:25]:
            if fn.startswith("."):
                continue
            out.append(f"{indent_files}{fn}")
    return "\n".join(out)


def get_deliverables_tree(input_dir: str) -> str:
    readme_path = _find_first_existing(input_dir, ["Readme.txt", "README.txt", "readme.txt"])
    if readme_path:
        tree = _extract_tree_from_readme(readme_path)
        if tree:
            return tree
    return _build_deliverables_tree_from_fs(input_dir)


def _as_float(v):
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    try:
        s = str(v).strip().replace(",", "")
        if s == "" or s.lower() in ("na", "nan", "none"):
            return None
        return float(s)
    except Exception:
        return None


def _downsample(points, max_points: int = 8000):
    if not points or len(points) <= max_points:
        return points
    step = max(1, len(points) // max_points)
    return points[::step]


def _norm_col(c: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(c).strip().lower())


def _find_col(df: pd.DataFrame, candidates) -> str | None:
    if df is None or df.empty:
        return None
    norm_map = {_norm_col(c): c for c in df.columns}
    for cand in candidates:
        nc = _norm_col(cand)
        if nc in norm_map:
            return norm_map[nc]
    return None


def _ensure_rows_have_keys(rows, keys, warning_message: str, extra: dict | None = None):
    if not rows:
        r = {k: "" for k in keys}
        first = keys[0] if keys else None
        if first:
            r[first] = warning_message
        if extra:
            r.update(extra)
        return [r]

    bad = False
    for r in rows:
        if not isinstance(r, dict):
            bad = True
            break
        for k in keys:
            if k not in r:
                bad = True
                break
        if bad:
            break

    if not bad:
        return rows

    r = {k: "" for k in keys}
    first = keys[0] if keys else None
    if first:
        r[first] = warning_message
    if extra:
        r.update(extra)
    return [r]


def _compute_dge_plot_data(df: pd.DataFrame) -> dict:
    import math

    fc_col = (
        _find_col(df, ["logfc", "log2fc", "log2foldchange", "logfoldchange", "lfc"])
        or next((c for c in df.columns if "logfc" in str(c).lower()), None)
        or next((c for c in df.columns if "log2fold" in str(c).lower()), None)
    )

    cpm_col = _find_col(df, ["logcpm", "log_cpm"]) or next((c for c in df.columns if "logcpm" in str(c).lower()), None)
    mean_col = _find_col(df, ["basemean", "aveexpr"]) or next((c for c in df.columns if "basemean" in str(c).lower()), None)

    sig_col = (
        _find_col(df, [
            "padj", "adjp", "adjpval", "fdr", "qvalue", "qval", "q_value",
            "pvalue", "pval", "p_value"
        ])
        or next((c for c in df.columns if any(tok in _norm_col(c) for tok in ("padj", "fdr", "qvalue", "qval"))), None)
        or next((c for c in df.columns if any(tok in _norm_col(c) for tok in ("pvalue", "pval"))), None)
    )

    ma = {"up": [], "down": [], "ns": []}
    volcano = {"up": [], "down": [], "ns": []}

    if not fc_col:
        return {"ma": ma, "volcano": volcano}

    for _, row in df.iterrows():
        fc = _as_float(row.get(fc_col))
        if fc is None:
            continue

        cpm = _as_float(row.get(cpm_col)) if cpm_col else None
        mean_expr = _as_float(row.get(mean_col)) if mean_col else None
        pval = _as_float(row.get(sig_col)) if sig_col else None

        is_sig = (pval is not None and pval > 0 and pval < 0.05)
        if is_sig and fc >= 1:
            bucket = "up"
        elif is_sig and fc <= -1:
            bucket = "down"
        else:
            bucket = "ns"

        ma_x = None
        if cpm is not None:
            ma_x = cpm
        elif mean_expr is not None and mean_expr >= 0:
            ma_x = math.log10(mean_expr + 1.0)
        if ma_x is not None:
            ma[bucket].append({"x": ma_x, "y": fc})

        if pval is not None and pval > 0:
            volcano[bucket].append({"x": fc, "y": -math.log10(pval)})

    for k in ma:
        ma[k] = _downsample(ma[k])
    for k in volcano:
        volcano[k] = _downsample(volcano[k])
    return {"ma": ma, "volcano": volcano}


def get_dge_figures(input_dir: str):
    dge_dir = os.path.join(input_dir, "05_differential_expression_analysis")
    if not os.path.isdir(dge_dir):
        return []

    figs = []
    for xlsx in sorted(glob.glob(os.path.join(dge_dir, "*DGE*.xlsx"))):
        base = os.path.splitext(os.path.basename(xlsx))[0]
        comp = re.sub(r"(?i)[_\-\s]*dge.*$", "", base).strip() or base

        heatmap = ""
        for pat in (
            f"*{comp}*Heatmap*.png",
            f"*{comp}*heatmap*.png",
            f"*{comp}*HeatMap*.png",
            f"*{comp}*HEATMAP*.png",
            f"*Heatmap*{comp}*.png",
            "*Heatmap*.png",
        ):
            matches = sorted(glob.glob(os.path.join(dge_dir, pat)))
            if matches:
                # Filter by comparison name if multiple heats exist
                match = next((m for m in matches if comp.lower() in os.path.basename(m).lower()), matches[0])
                heatmap = match
                break

        ma_vol_pdf = ""
        for pat in (
            f"*{comp}*MA*Volcano*.pdf",
            f"*{comp}*Volcano*.pdf",
            f"*{comp}*MA*.pdf",
            f"*MA*Volcano*{comp}*.pdf",
            "*Volcano*.pdf",
            "*MA*.pdf",
        ):
            matches = sorted(glob.glob(os.path.join(dge_dir, pat)))
            if matches:
                 match = next((m for m in matches if comp.lower() in os.path.basename(m).lower()), matches[0])
                 ma_vol_pdf = match
                 break

        plot_data = None
        try:
            df = pd.read_excel(xlsx)
            plot_data = _compute_dge_plot_data(df)
        except Exception:
            plot_data = {"ma": {"up": [], "down": [], "ns": []}, "volcano": {"up": [], "down": [], "ns": []}}

        fig = {
            "comparison": comp,
            "dge_xlsx_src": _relpath_posix(xlsx, input_dir),
            "heatmap_png_src": _relpath_posix(heatmap, input_dir) if os.path.exists(heatmap) else "",
            "ma_volcano_pdf_src": _relpath_posix(ma_vol_pdf, input_dir) if os.path.exists(ma_vol_pdf) else "",
            "plot_data_json": json.dumps(plot_data),
            "ma_plot_data": bool(plot_data and any(plot_data.get("ma", {}).get(k) for k in ("up", "down", "ns"))),
            "volcano_plot_data": bool(plot_data and any(plot_data.get("volcano", {}).get(k) for k in ("up", "down", "ns"))),
        }
        figs.append(fig)

    return figs


def _read_table_preview(path: str, n: int = 10):
    try:
        if path.lower().endswith((".xlsx", ".xls")):
            df = pd.read_excel(path)
        else:
            df = pd.read_csv(path)
        df = df.head(n)
        return df.fillna("").to_dict(orient="records")
    except Exception:
        return []


def get_functional_assets(input_dir: str):
    candidates = []
    for ext in ("*.xlsx", "*.xls", "*.csv", "*.tsv"):
        candidates.extend(glob.glob(os.path.join(input_dir, "**", ext), recursive=True))

    def _pick(pred):
        for p in candidates:
            name = os.path.basename(p).lower()
            if pred(name):
                return p
        return ""

    go_xlsx = _pick(lambda n: "go" in n and n.endswith((".xlsx", ".xls")))
    go_csv = _pick(lambda n: "go" in n and n.endswith((".csv", ".tsv")))
    kegg_xlsx = _pick(lambda n: "kegg" in n and n.endswith((".xlsx", ".xls")))
    kegg_csv = _pick(lambda n: "kegg" in n and n.endswith((".csv", ".tsv")))

    images = []
    for ext in ("*.png", "*.jpg", "*.jpeg", "*.svg"):
        images.extend(glob.glob(os.path.join(input_dir, "**", ext), recursive=True))

    kegg_path_img = next((p for p in images if "kegg" in os.path.basename(p).lower() and "path" in os.path.basename(p).lower()), "")
    enrich_img = next((p for p in images if "enrich" in os.path.basename(p).lower() or "dotplot" in os.path.basename(p).lower()), "")

    go_best = go_xlsx or go_csv
    kegg_best = kegg_xlsx or kegg_csv

    return {
        "go_results_xlsx": _relpath_posix(go_xlsx, input_dir) if go_xlsx else "",
        "kegg_results_xlsx": _relpath_posix(kegg_xlsx, input_dir) if kegg_xlsx else "",
        "go_results_preview": _read_table_preview(go_best) if go_best else [],
        "kegg_results_preview": _read_table_preview(kegg_best) if kegg_best else [],
        "kegg_pathway_image_src": _relpath_posix(kegg_path_img, input_dir) if kegg_path_img else "",
        "enrichment_image_src": _relpath_posix(enrich_img, input_dir) if enrich_img else "",
    }


def _find_component_asset(script_dir: str, filename: str) -> str:
    path = os.path.join(script_dir, "components", filename)
    return path if os.path.exists(path) else ""


def _find_first_component_asset(script_dir: str, filenames) -> str:
    for fn in filenames:
        p = _find_component_asset(script_dir, fn)
        if p:
            return p
    return ""


def _load_metadata_json(input_dir: str, metadata_path: str = None) -> dict:
    """Load per-project metadata.json.

    Supports either:
    - explicit metadata_path
    - input_dir/metadata.json (default)
    """
    try:
        path = metadata_path or os.path.join(input_dir, "metadata.json")
        if not path or not os.path.exists(path):
            return {}
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _compute_default_output_path(input_dir: str, project_id: str, output_arg: str) -> str:
    """Keep CLI backward compatibility while defaulting output into project root."""
    if not output_arg or output_arg.strip() == "":
        return os.path.join(input_dir, f"{project_id}_report.html")

    # If user left the default "report.html", put it under project root with a clearer name
    if os.path.basename(output_arg) == "report.html" and os.path.dirname(output_arg) in ("", "."):
        return os.path.join(input_dir, f"{project_id}_report.html")

    # If user passed a relative path, interpret it relative to input_dir
    if not os.path.isabs(output_arg):
        return os.path.join(input_dir, output_arg)

    return output_arg

def load_static_content(script_dir):
    static_content_path = os.path.join(script_dir, "report_static_content.json")
    if os.path.exists(static_content_path):
        with open(static_content_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def _choose_file(files):
    if not files:
        return None
    existing = [f for f in files if os.path.exists(f)]
    if not existing:
        return None
    existing.sort(key=lambda p: (os.path.getmtime(p), os.path.getsize(p)), reverse=True)
    return existing[0]


def _find_first_existing(input_dir, relative_globs):
    candidates = []
    for rel_glob in relative_globs:
        candidates.extend(glob.glob(os.path.join(input_dir, rel_glob)))
    return _choose_file(candidates)


def _parse_project_readme(input_dir):
    readme_path = _find_first_existing(input_dir, ["Readme.txt", "README.txt", "readme.txt"])
    details = {}
    if not readme_path:
        return details

    try:
        with open(readme_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        patterns = {
            "project_id": r"Project\s*ID\s*:\s*(.+)",
            "project_pi": r"Project\s*PI\s*:\s*(.+)",
            "application": r"Application\s*:\s*(.+)",
            "no_of_samples": r"No\s*of\s*Samples\s*:\s*(.+)",
        }
        for key, pat in patterns.items():
            m = re.search(pat, content, flags=re.IGNORECASE)
            if m:
                details[key] = m.group(1).strip()
    except Exception:
        return details

    return details


def _validate_deliverables_structure(input_dir):
    warnings = []
    required_dirs = [
        "01_Raw_Data",
        "02_reference_genome_and_gff",
        "03_transcript_assembly_gtf",
        "04_transcript_sequences_fasta",
        "05_differential_expression_analysis",
    ]
    optional_dirs = [
        "06_Significant_DGE_GO",
        "07_Significant_DGE_pathways",
        "08_Significant_DGE_Enrichment",
    ]

    for d in required_dirs:
        if not os.path.isdir(os.path.join(input_dir, d)):
            warnings.append(f"Missing required directory: {d}")
    for d in optional_dirs:
        if not os.path.isdir(os.path.join(input_dir, d)):
            warnings.append(f"Missing optional directory: {d}")

    raw_dir = os.path.join(input_dir, "01_Raw_Data")
    if os.path.isdir(raw_dir):
        fastqs = glob.glob(os.path.join(raw_dir, "*.fastq.gz"))
        if not fastqs:
            warnings.append("No .fastq.gz files found in 01_Raw_Data (sample detection may rely on Mapping/Stats).")

    return warnings

def parse_arguments():
    parser = argparse.ArgumentParser(description="Generate NGS Report")
    parser.add_argument("--input", required=True, help="Path to the project directory")
    parser.add_argument("--output", default="report.html", help="Output HTML file path")
    parser.add_argument("--strict", action="store_true", help="Fail if required inputs are missing")
    parser.add_argument("--client-name", default=None, help="Client name to display on the cover page")
    parser.add_argument("--client-org", default=None, help="Client organization to display on the cover page")
    parser.add_argument("--reference-organism", default=None, help="Reference organism name used in the report")
    parser.add_argument("--logo-path", default=None, help="Logo path to use in the report")
    parser.add_argument("--metadata-path", default=None, help="Optional path to metadata.json (defaults to <input>/metadata.json)")
    return parser.parse_args()

def get_project_details(input_dir, metadata_path: str = None):
    # Try to infer project ID from directory name or a config
    project_id = os.path.basename(os.path.normpath(input_dir))
    parsed = _parse_project_readme(input_dir)
    meta = _load_metadata_json(input_dir, metadata_path)

    # Prefer metadata.json over Readme.txt
    if meta.get("project_id"):
        project_id = str(meta.get("project_id")).strip()
    if parsed.get("project_id"):
        project_id = parsed["project_id"]

    client_name = meta.get("client_name") or "Client Name"
    client_org = meta.get("client_org") or "Client Organization"
    return {
        "project_id": project_id,
        "report_date": datetime.now().strftime("%d-%m-%Y"),
        "client_name": client_name,
        "client_org": client_org,
        "project_pi": (meta.get("pi_name") or meta.get("project_pi") or parsed.get("project_pi", "")),
        "application": (meta.get("application") or parsed.get("application", "")),
        "genome": (meta.get("genome") or meta.get("reference_genome") or ""),
        "no_of_samples": (meta.get("no_of_samples") or parsed.get("no_of_samples", "")),
    }

def get_samples(input_dir):
    raw_data_dir = os.path.join(input_dir, "01_Raw_Data")
    if not os.path.exists(raw_data_dir):
        return []
    
    samples = set()
    
    # Priority 1: Try to look for fastq files
    files = glob.glob(os.path.join(raw_data_dir, "*.fastq.gz"))
    for f in files:
        basename = os.path.basename(f)
        match = re.match(r"(.*)_R[12].*\.fastq\.gz", basename)
        if match:
            samples.add(match.group(1))

    # Priority 2: If no fastq, extract from Mapping file (txt or xlsx)
    if not samples:
        mapping_files = glob.glob(os.path.join(raw_data_dir, "Mapping.*"))
        mapping_files = sorted(mapping_files)
        for mapping_file in mapping_files:
            try:
                if mapping_file.endswith('.xlsx'):
                    df = pd.read_excel(mapping_file)
                else:
                    df = pd.read_csv(mapping_file, sep=None, engine='python')
                
                df.columns = df.columns.str.strip()
                sample_col = next((c for c in df.columns if 'Sample' in c), None)
                if sample_col:
                    samples.update(df[sample_col].dropna().astype(str).unique())
            except Exception as e:
                print(f"Error reading {mapping_file} in get_samples: {e}")
    
    if not samples:
        stats_files = glob.glob(os.path.join(raw_data_dir, "*Raw*Stats*"))
        # Also check for Rawstats.xlsx specifically if glob didn't catch it
        if not stats_files:
            stats_files = glob.glob(os.path.join(raw_data_dir, "*Rawstats*"))
            
        for f in stats_files:
            try:
                if f.endswith('.xlsx'):
                    df = pd.read_excel(f)
                else:
                    df = pd.read_csv(f, sep="\t")
                
                df.columns = df.columns.str.strip()
                sample_col = next((c for c in df.columns if 'Sample' in c), None) # 'Sample' or 'SampleID'
                if sample_col:
                    samples.update(df[sample_col].dropna().astype(str).unique())
            except Exception as e:
                print(f"Error reading Stats file {f} in get_samples: {e}")

    return sorted(list(samples))

def get_sequencing_stats(input_dir, samples):
    # Parse *Raw_Stats_with_GB.txt or .xlsx in 01_Raw_Data
    stats = []
    raw_data_dir = os.path.join(input_dir, "01_Raw_Data")
    
    # Try txt then fallback to xlsx or wildcard
    stats_files = glob.glob(os.path.join(raw_data_dir, "*Raw*Stats*"))
    if not stats_files:
        stats_files = glob.glob(os.path.join(raw_data_dir, "*Rawstats*"))
    stats_files = sorted(stats_files)
    
    if stats_files:
        try:
            f = _choose_file(stats_files)
            if f.endswith('.xlsx'):
                df = pd.read_excel(f)
            else:
                df = pd.read_csv(f, sep="\t")

            # Normalize columns
            df.columns = df.columns.str.strip()
            
            for _, row in df.iterrows():
                # Flexible matching
                sample_col = next((c for c in df.columns if 'sample' in c.lower()), None)
                # Specific columns from user's file: "Total Reads(R1+R2)", "Total Bases (R1+R2)", "Total Data(GB)"
                reads_col = next((c for c in df.columns if '(R1+R2)' in c and 'Reads' in c), None) 
                if not reads_col: reads_col = next((c for c in df.columns if 'ReadCount' in c), None)
                
                bases_col = next((c for c in df.columns if '(R1+R2)' in c and 'Bases' in c), None)
                if not bases_col: bases_col = next((c for c in df.columns if 'BaseCount' in c), None)

                gb_col = next((c for c in df.columns if 'Total Data' in c), None)

                if sample_col:
                    reads_val = str(row[reads_col]).replace(',', '') if reads_col and pd.notna(row[reads_col]) else "0"
                    bases_val = str(row[bases_col]).replace(',', '') if bases_col and pd.notna(row[bases_col]) else "0"
                    gb_val = str(row[gb_col]).replace(',', '') if gb_col and pd.notna(row[gb_col]) else "0"

                    try:
                        reads_fmt = f"{int(float(reads_val)):,}"
                    except: reads_fmt = "N/A"
                    
                    try:
                        bases_fmt = f"{int(float(bases_val)):,}"
                    except: bases_fmt = "N/A"

                    stats.append({
                        "sample": str(row[sample_col]),
                        "total_reads": reads_fmt,
                        "total_bases": bases_fmt,
                        "data_gb": gb_val
                    })
        except Exception as e:
            print(f"Error parsing raw stats file: {e}")
    
    # If no stats found
    if not stats:
        for s in samples:
            stats.append({
                "sample": s,
                "total_reads": "N/A",
                "total_bases": "N/A",
                "data_gb": "N/A"
            })
    return stats

def get_mapping_stats(input_dir):
    raw_data_dir = os.path.join(input_dir, "01_Raw_Data")
    mapping_stats = []
    mapping_files = glob.glob(os.path.join(raw_data_dir, "Mapping.*"))
    mapping_files = sorted(mapping_files)
    
    if mapping_files:
        try:
            f = _choose_file(mapping_files)
            # Fix: sep=None with engine='python' allows sniffing, 
            # and avoid passing regex characters that conflict
            if f.endswith('.xlsx'):
                df = pd.read_excel(f)
            else:
                df = pd.read_csv(f, sep=r'\s+|\t', engine='python')
                
            df.columns = df.columns.str.strip()

            for _, row in df.iterrows():
                # Columns: Sample Name, Total Reads, No. of Mapped reads, % of mapped reads, # uniquely mapped reads, % uniquely mapped reads
                sample_col = next((c for c in df.columns if 'Sample' in c), None)
                mapped_col = _find_col(df, ["mappedreads", "no_of_mapped_reads"]) or next((c for c in df.columns if 'Mapped' in c and 'reads' in c.lower() and '%' not in c), None)
                pct_mapped_col = _find_col(df, ["percentmapped", "pct_mapped_reads"]) or next((c for c in df.columns if '%' in c and 'mapped' in c.lower()), None)
                unique_col = _find_col(df, ["uniquelymappedreads", "no_of_uniquely_mapped_reads"]) or next((c for c in df.columns if 'uniquely' in c.lower() and 'mapped' in c.lower() and '%' not in c), None)
                pct_unique_col = _find_col(df, ["percentuniquelymapped", "pct_uniquely_mapped_reads"]) or next((c for c in df.columns if '%' in c and 'uniquely' in c.lower()), None)
                total_reads_col = _find_col(df, ["totalreads", "no_of_reads"]) or next((c for c in df.columns if 'Total' in c and 'Reads' in c), None)

                if sample_col:
                     mapping_stats.append({
                        "sample_name": row[sample_col],
                        "total_reads": f"{int(float(str(row[total_reads_col]).replace(',',''))):,}" if total_reads_col and pd.notna(row[total_reads_col]) else "N/A",
                        "mapped_reads": f"{int(float(str(row[mapped_col]).replace(',',''))):,}" if mapped_col and pd.notna(row[mapped_col]) else "N/A",
                        "percent_mapped": str(row[pct_mapped_col]) if pct_mapped_col and pd.notna(row[pct_mapped_col]) else "N/A",
                        "unique_reads": f"{int(float(str(row[unique_col]).replace(',',''))):,}" if unique_col and pd.notna(row[unique_col]) else "N/A",
                        "percent_unique": str(row[pct_unique_col]) if pct_unique_col and pd.notna(row[pct_unique_col]) else "N/A"
                    })
        except Exception as e:
            print(f"Error parsing Mapping.txt: {e}")

    # Fallback: STAR Log.final.out parsing
    if not mapping_stats and os.path.exists(raw_data_dir):
        log_files = glob.glob(os.path.join(raw_data_dir, "**", "*Log.final.out"), recursive=True)
        log_files = sorted(log_files)

        def _parse_star_log(path: str) -> dict:
            vals = {}
            with open(path, "r", encoding="utf-8", errors="ignore") as fh:
                for line in fh:
                    if "|" not in line:
                        continue
                    k, v = line.split("|", 1)
                    vals[k.strip()] = v.strip()

            total_reads = vals.get("Number of input reads")
            uniq_pct = vals.get("Uniquely mapped reads %")
            multi_pct = vals.get("% of reads mapped to multiple loci") or vals.get("% of reads mapped to multiple loci")
            mapped_pct = None
            try:
                u = float(str(uniq_pct).replace("%", "")) if uniq_pct is not None else 0.0
                m = float(str(multi_pct).replace("%", "")) if multi_pct is not None else 0.0
                mapped_pct = u + m
            except Exception:
                mapped_pct = None

            return {
                "total_reads": total_reads,
                "percent_unique": uniq_pct,
                "percent_mapped": f"{mapped_pct:.2f}%" if mapped_pct is not None else "N/A",
            }

        for lf in log_files:
            sample_name = os.path.basename(os.path.dirname(lf))
            parsed = _parse_star_log(lf)
            mapping_stats.append({
                "sample_name": sample_name,
                "total_reads": parsed.get("total_reads") or "N/A",
                "mapped_reads": "N/A",
                "percent_mapped": parsed.get("percent_mapped") or "N/A",
                "unique_reads": "N/A",
                "percent_unique": parsed.get("percent_unique") or "N/A",
            })

    if not mapping_stats:
        mapping_stats = _ensure_rows_have_keys(
            [],
            ["sample_name", "total_reads", "mapped_reads", "percent_mapped", "unique_reads", "percent_unique"],
            "Mapping stats unavailable",
        )
    return mapping_stats

def get_ref_stats(input_dir):
    ref_dir = os.path.join(input_dir, "02_reference_genome_and_gff")
    stats = {
        "total_scaffolds": 0,
        "genome_length": 0,
        "mean_scaffold_size": 0,
        "max_scaffold_size": 0
    }
    total_genes = 0
    
    if os.path.exists(ref_dir):
        # Parse Fasta for genome stats
        fasta_files = glob.glob(os.path.join(ref_dir, "*.fa")) + glob.glob(os.path.join(ref_dir, "*.fasta"))
        if fasta_files:
            scaffold_lengths = []
            fasta_path = _choose_file(fasta_files)
            with open(fasta_path, 'r') as f:
                curr_len = 0
                for line in f:
                    if line.startswith(">"):
                        if curr_len > 0:
                            scaffold_lengths.append(curr_len)
                        curr_len = 0
                    else:
                        curr_len += len(line.strip())
                if curr_len > 0:
                    scaffold_lengths.append(curr_len)
            
            if scaffold_lengths:
                stats["total_scaffolds"] = len(scaffold_lengths)
                stats["genome_length"] = f"{sum(scaffold_lengths):,}"
                stats["mean_scaffold_size"] = f"{int(sum(scaffold_lengths)/len(scaffold_lengths)):,}"
                stats["max_scaffold_size"] = f"{max(scaffold_lengths):,}"

        # Parse GFF for gene count (3rd column feature type)
        gff_files = glob.glob(os.path.join(ref_dir, "*.gff")) + glob.glob(os.path.join(ref_dir, "*.gff3"))
        if gff_files:
            gff_path = _choose_file(gff_files)
            with open(gff_path, 'r', encoding="utf-8", errors="ignore") as f:
                for line in f:
                    if not line or line.startswith("#"):
                        continue
                    parts = line.rstrip("\n").split("\t")
                    if len(parts) < 3:
                        continue
                    feature = parts[2].strip() # Case-insensitive comparison below
                    if feature.lower() == "gene":
                        total_genes += 1
    
    return stats, f"{total_genes:,}"


def get_dge_comparison_tables(input_dir: str):
    dge_dir = os.path.join(input_dir, "05_differential_expression_analysis")
    if not os.path.isdir(dge_dir):
        return (
            {"headers": ["Comparison", "Source"], "rows": [["Comparison table unavailable", ""]]},
            {"headers": ["Group", "Samples"], "rows": [["Group table unavailable", ""]]},
        )

    files = sorted(glob.glob(os.path.join(dge_dir, "*DGE*.xlsx")))
    if not files:
        return (
            {"headers": ["Comparison", "Source"], "rows": [["No DGE files found", ""]]},
            {"headers": ["Group", "Samples"], "rows": [["Group table unavailable", ""]]},
        )

    comp_rows = []
    groups = {}

    for f in files:
        base = os.path.basename(f)
        comp = re.sub(r"(?i)[_\-\s]*dge.*\.xlsx$", "", base).strip() or base
        comp_rows.append([comp, os.path.basename(f)])

        tokens = re.split(r"(?i)\bvs\b|[_\-]+", comp)
        tokens = [t.strip() for t in tokens if t and t.strip()]
        if len(tokens) >= 2:
            groups.setdefault(tokens[0], []).append(comp)
            groups.setdefault(tokens[1], []).append(comp)

    dge_comparison_table = {"headers": ["Comparison", "Source"], "rows": comp_rows or [["Comparison table unavailable", ""]]}

    group_rows = []
    for g, comps in sorted(groups.items()):
        group_rows.append([g, ", ".join(sorted(set(comps)))])
    dge_group_table = {"headers": ["Group", "Comparisons"], "rows": group_rows or [["Group table unavailable", ""]]}

    return dge_comparison_table, dge_group_table

def get_assembly_stats(input_dir, samples):
    assembly_stats = []
    # Logic: Read 04_transcript_sequences_fasta for merged and sample-wise fastas
    fasta_dir = os.path.join(input_dir, "04_transcript_sequences_fasta")
    
    def calc_fasta_stats(filepath):
        lengths = []
        with open(filepath, 'r') as f:
            curr = 0
            for line in f:
                if line.startswith(">"):
                    if curr > 0: lengths.append(curr)
                    curr = 0
                else:
                    curr += len(line.strip())
            if curr > 0: lengths.append(curr)
        
        if not lengths: return 0, 0, 0
        return len(lengths), sum(lengths), int(sum(lengths)/len(lengths))

    if os.path.exists(fasta_dir):
        # Merged
        merged_files = glob.glob(os.path.join(fasta_dir, "Merged.fasta"))
        if merged_files:
             merged_path = _choose_file(merged_files)
             count, total, mean = calc_fasta_stats(merged_path)
             assembly_stats.append({
                 "sample": "merged.fasta",
                 "num_transcripts": f"{count:,}",
                 "total_bp": f"{total:,}",
                 "mean_size": f"{mean:,}"
             })
        
        # Per Sample (trying to match sample names)
        for s in samples:
             # Try to find a file that looks like it belongs to the sample
             # Pattern from user: NGS_240540_PM-CONTROL-2Aligned_transcript.fasta
             # It's fuzzy, so we look for files containing the sample name
             files = glob.glob(os.path.join(fasta_dir, f"*{s}*.fasta"))
             for f in files:
                if "Merged" in f: continue
                count, total, mean = calc_fasta_stats(f)
                assembly_stats.append({
                 "sample": os.path.basename(f),
                 "num_transcripts": f"{count:,}",
                 "total_bp": f"{total:,}",
                 "mean_size": f"{mean:,}"
             })

    return assembly_stats

def get_dge_stats(input_dir):
    dge_dir = os.path.join(input_dir, "05_differential_expression_analysis")
    stats = []
    labels = []
    up_counts = []
    down_counts = []
    
    if os.path.exists(dge_dir):
        files = glob.glob(os.path.join(dge_dir, "*DGE*.xlsx"))
        for f in files:
            try:
                df = pd.read_excel(f)
                # Assuming standard edgeR output or similar with logFC and FDR/PValue
                # Adjust column names as per actual files
                # User's file implies: Comparison1_DGE.xlsx
                comparison_name = re.sub(r"(?i)[_\-\s]*dge.*\.xlsx$", "", os.path.basename(f)).strip()

                fc_col = (
                    _find_col(df, ["logfc", "log2fc", "log2foldchange", "logfoldchange", "lfc"])
                    or next((c for c in df.columns if "logfc" in str(c).lower()), None)
                    or next((c for c in df.columns if "log2fold" in str(c).lower()), None)
                )

                p_col = (
                    _find_col(df, ["padj", "adjp", "adjpval", "fdr", "qvalue", "qval", "pvalue", "pval"])
                    or next((c for c in df.columns if any(tok in _norm_col(c) for tok in ("padj", "fdr", "qvalue", "qval", "pvalue", "pval"))), None)
                )
                
                if fc_col and p_col:
                    sig_up = len(df[(df[fc_col] > 1) & (df[p_col] < 0.05)])
                    sig_down = len(df[(df[fc_col] < -1) & (df[p_col] < 0.05)])
                    total_sig = sig_up + sig_down
                    
                    stats.append({
                        "comparison": comparison_name,
                        "total_degs": len(df), # Or maybe total genes tested? usually DGE file has all genes
                        "sig_down": sig_down,
                        "sig_up": sig_up,
                        "total_sig": total_sig
                    })
                    
                    labels.append(comparison_name)
                    up_counts.append(sig_up)
                    down_counts.append(sig_down)
            except Exception as e:
                print(f"Error reading DGE file {f}: {e}")

    return stats, labels, up_counts, down_counts

def get_pathway_stats(input_dir):
    # This usually parses the KEGG/GO result files
    # User's path: 07_Significant_DGE_pathways/Comparison1_Significant_DGE_Pathways.xlsx
    path_dir = os.path.join(input_dir, "07_Significant_DGE_pathways")
    stats = []
    files = []
    if os.path.exists(path_dir):
        files = sorted(glob.glob(os.path.join(path_dir, "*.xlsx")))

    # Fallback: many deliveries store pathway tables elsewhere (e.g. 05/06/08/files)
    if not files:
        files = sorted(glob.glob(os.path.join(input_dir, "**", "*Significant*DGE*Pathway*.xlsx"), recursive=True))
    if not files:
        files = sorted(glob.glob(os.path.join(input_dir, "**", "*Significant*DGE*Pathways*.xlsx"), recursive=True))
    if not files:
        files = sorted(glob.glob(os.path.join(input_dir, "**", "*DGE*Pathway*.xlsx"), recursive=True))

    if files:
        try:
            f = _choose_file(files)
            df = pd.read_excel(f)
            df.columns = df.columns.str.strip()

            level1_col = (
                _find_col(df, ["level1", "level1category", "category", "class", "keggclass"])
                or next((c for c in df.columns if any(tok in _norm_col(c) for tok in ("level1", "category", "class"))), None)
            )
            level2_col = (
                _find_col(df, ["level2", "level2subcategory", "subcategory", "pathway", "keggpathway"])
                or next((c for c in df.columns if any(tok in _norm_col(c) for tok in ("level2", "subcategory", "pathway"))), None)
            )
            count_col = (
                _find_col(df, ["count", "counts", "gene", "genes", "gene_count", "genecount"])
                or next((c for c in df.columns if any(tok in _norm_col(c) for tok in ("count", "gene"))), None)
            )

            if not (level1_col and level2_col and count_col):
                return _ensure_rows_have_keys(
                    [],
                    ["level1", "level2", "count"],
                    "Pathway columns not recognized",
                    {"level2": f"File: {_relpath_posix(f, input_dir)}"},
                )

            for _, row in df.head(50).iterrows():
                stats.append({
                    "level1": str(row.get(level1_col, "")).strip(),
                    "level2": str(row.get(level2_col, "")).strip(),
                    "count": int(_as_float(row.get(count_col, 0)) or 0)
                })
        except Exception as e:
            print(f"Error reading pathway file: {e}")

    if not files:
        return _ensure_rows_have_keys(
            [],
            ["level1", "level2", "count"],
            "No pathway file found",
            {"level2": "Expected in 07_Significant_DGE_pathways or matching *Pathway*.xlsx"},
        )

    # Final validation: if template expects level1 and it's missing from ALL rows, emit a visible warning row
    if stats and not any(s.get("level1") for s in stats):
         return _ensure_rows_have_keys(
            [],
            ["level1", "level2", "count"],
            "Pathway columns not recognized",
            {"level2": "Found file but 'level1' column was missing or empty."}
        )

    return _ensure_rows_have_keys(stats, ["level1", "level2", "count"], "Pathway stats unavailable")

def get_go_distribution(input_dir):
    """Parses GO distribution stats from 06_Significant_DGE_GO or similar."""
    go_dir = os.path.join(input_dir, "06_Significant_DGE_GO")
    stats = []
    
    # Prioritize specific filenames
    candidates = []
    if os.path.exists(go_dir):
        candidates.extend(glob.glob(os.path.join(go_dir, "*with_GO*.xlsx")))
        candidates.extend(glob.glob(os.path.join(go_dir, "*GO*Distribution*.xlsx")))
        candidates.extend(glob.glob(os.path.join(go_dir, "*GO*results*.xlsx")))
        # Fallback to any xlsx that HAS 'GO' in name
        if not candidates:
            candidates.extend(glob.glob(os.path.join(go_dir, "*GO*.xlsx")))

    if not candidates:
        candidates.extend(glob.glob(os.path.join(input_dir, "**", "*GO*Distribution*.xlsx"), recursive=True))
        candidates.extend(glob.glob(os.path.join(input_dir, "**", "*with_GO.xlsx"), recursive=True))

    if candidates:
        try:
            f = _choose_file(candidates)
            df = pd.read_excel(f)
            df.columns = df.columns.str.strip()
            
            # Case 1: Pre-summarized table
            term_col = _find_col(df, ["term", "go_term", "description", "goterm"]) 
            count_col = _find_col(df, ["count", "gene_count", "genes"])
            ont_col = _find_col(df, ["ontology", "category", "namespace", "class"])

            if term_col and count_col:
                for _, row in df.head(30).iterrows():
                    stats.append({
                        "term": str(row[term_col]),
                        "count": int(_as_float(row[count_col]) or 0),
                        "ontology": str(row[ont_col]) if ont_col else "N/A"
                    })
                return stats

            # Case 2: Raw annotation file (e.g. Comparison1_Significant_DGE_with_GO.xlsx)
            # We want to aggregate counts by Category and Term
            cat_col = (
                _find_col(df, ["annotationgocategory", "annotation_go_category", "go_category", "category", "ontology"]) 
                or next((c for c in df.columns if any(t in _norm_col(c) for t in ("annotationgocategory", "gocategory", "ontol"))), None)
            )
            term_col = (
                _find_col(df, ["annotationgoterm", "annotation_go_term", "go_term", "term"])
                or next((c for c in df.columns if any(t in _norm_col(c) for t in ("annotationgoterm", "goterm"))), None)
            )
            
            if cat_col and term_col:
                counts = {}
                for _, row in df.iterrows():
                    c_val = str(row[cat_col]) if pd.notna(row[cat_col]) else ""
                    t_val = str(row[term_col]) if pd.notna(row[term_col]) else ""
                    cats = c_val.split(';')
                    terms = t_val.split(';')
                    for c, t in zip(cats, terms):
                        c, t = c.strip(), t.strip()
                        if c and t:
                            key = (c, t)
                            counts[key] = counts.get(key, 0) + 1
                
                if counts:
                    sorted_counts = sorted(counts.items(), key=lambda x: x[1], reverse=True)
                    for (cat, term), count in sorted_counts[:30]:
                        stats.append({
                            "ontology": cat,
                            "term": term,
                            "count": count
                        })
                    return stats

        except Exception as e:
            print(f"Error parsing GO distribution: {e}")

    return _ensure_rows_have_keys(stats, ["term", "count", "ontology"], "GO distribution unavailable")

def main():
    args = parse_arguments()
    input_dir = args.input
    output_file = args.output

    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    print(f"Processing project: {input_dir}")
    
    warnings = _validate_deliverables_structure(input_dir)
    if args.strict and any(w.startswith("Missing required directory") for w in warnings):
        raise SystemExit("\n".join(warnings))

    # 1. Get Project Details (metadata.json -> Readme.txt -> defaults)
    meta = _load_metadata_json(input_dir, args.metadata_path)
    project_details = get_project_details(input_dir, metadata_path=args.metadata_path)

    # Compute output path (default into project root)
    output_file = _compute_default_output_path(input_dir, project_details.get("project_id") or "report", output_file)
    
    # 2. Get Samples
    samples = get_samples(input_dir)
    print(f"Found {len(samples)} samples: {samples}")
    
    # 3. Get Sequencing Stats
    seq_stats = get_sequencing_stats(input_dir, samples)
    
    # 4. Reference Genome Stats
    ref_stats, total_genes = get_ref_stats(input_dir)
    
    # 5. Mapping Stats
    mapping_stats = get_mapping_stats(input_dir)
    # Fallback if empty logic (optional)
    if not mapping_stats:
        # Try to infer or leave empty
        for s in samples:
            mapping_stats.append({
                "sample_name": s,
                "mapped_reads": "N/A",
                "percent_mapped": "N/A",
                "unique_reads": "N/A",
                "percent_unique": "N/A"
            })

    # 6. Assembly Stats
    assembly_stats = get_assembly_stats(input_dir, samples)
    
    # 7. DGE Stats
    dge_stats, dge_labels, dge_up, dge_down = get_dge_stats(input_dir)
    
    # 8. Component Stats (Pathway)
    pathway_stats = get_pathway_stats(input_dir)
    go_distribution = get_go_distribution(input_dir)

    # 9. DGE Figures (MA/Volcano interactive + heatmaps)
    dge_figures = get_dge_figures(input_dir)

    # 9b. DGE comparison/group tables (data-driven)
    dge_comparison_table, dge_group_table = get_dge_comparison_tables(input_dir)

    # 10. Functional assets (GO/KEGG enrichment previews)
    func_assets = get_functional_assets(input_dir)

    # 11. Deliverables tree
    deliverables_tree = get_deliverables_tree(input_dir)

    # Static component assets (repo-local)
    gffcompare_codes_path = _find_component_asset(script_dir, "gffcompare_codes.png")
    gffcompare_codes_src = _relpath_posix(gffcompare_codes_path, input_dir) if gffcompare_codes_path else ""

    workflow_figure_path = _find_first_component_asset(script_dir, ["workflow.png", "workflow.svg"])
    workflow_figure_src = _relpath_posix(workflow_figure_path, input_dir) if workflow_figure_path else ""

    stringtie_merge_path = _find_component_asset(script_dir, "stringtie_merge_illustration.svg")
    stringtie_merge_figure_src = _relpath_posix(stringtie_merge_path, input_dir) if stringtie_merge_path else ""

    isoforms_path = _find_component_asset(script_dir, "isoforms.png")
    isoforms_figure_src = _relpath_posix(isoforms_path, input_dir) if isoforms_path else ""

    pathway_ex_path = _find_component_asset(script_dir, "pathway_ex.png")
    pathway_ex_figure_src = _relpath_posix(pathway_ex_path, input_dir) if pathway_ex_path else ""

    client_name = args.client_name or project_details.get('client_name')
    client_org = args.client_org or project_details.get('client_org')
    reference_organism = (
        args.reference_organism
        if args.reference_organism is not None
        else (meta.get("reference_organism") or meta.get("organism") or "Organism Name")
    )
    # Prefer explicit CLI --logo-path; else metadata.json can set logo_path; else default asset
    logo_path = args.logo_path or meta.get("logo_path") or os.path.join(script_dir, 'assets', 'logo.png')

    # Render
    static_content = load_static_content(script_dir)

    # Prepare Context
    context = {
        "project_id": project_details["project_id"],
        "report_date": project_details["report_date"],
        "client_name": client_name,
        "client_org": client_org,
        "project_pi": project_details.get("project_pi", ""),
        "application": project_details.get("application", ""),
        "no_of_samples": project_details.get("no_of_samples", ""),
        "sample_count": len(samples),
        "samples": samples,
        "qubit_data": [], # Placeholder
        "library_sizes": [450] * len(samples), # Placeholder
        "sequencing_stats": seq_stats,
        "reference_organism": reference_organism,
        "total_genes": total_genes,
        "ref_stats": ref_stats,
        "mapping_stats": mapping_stats,
        "total_transcripts": sum(int(x['num_transcripts'].replace(',', '')) for x in assembly_stats if x['sample'] == 'merged.fasta') if assembly_stats else 0,
        "mean_transcript_size": next((x['mean_size'] for x in assembly_stats if x['sample'] == 'merged.fasta'), 0),
        "assembly_stats": assembly_stats,
        "diff_expr_stats": dge_stats,
        "dge_chart_labels": dge_labels,
        "dge_chart_up": dge_up,
        "dge_chart_down": dge_down,
        "pathway_stats": pathway_stats,
        "go_distribution": go_distribution,
        "pathway_image_src": "", # Logic to find image if exists
        "dge_figures": dge_figures,
        "dge_comparison_table": dge_comparison_table,
        "dge_group_table": dge_group_table,
        "func_assets": func_assets,
        "deliverables_tree": deliverables_tree,
        "gffcompare_codes_src": gffcompare_codes_src,
        "workflow_figure_src": workflow_figure_src,
        "stringtie_merge_figure_src": stringtie_merge_figure_src,
        "isoforms_figure_src": isoforms_figure_src,
        "pathway_ex_figure_src": pathway_ex_figure_src,
        "logo_path": logo_path,
        "warnings": warnings,
        "static_content": static_content
    }
    
    templates_dir = os.path.join(script_dir, 'templates')
    env = Environment(loader=FileSystemLoader(templates_dir))
    template = env.get_template('report_template.html')
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(template.render(context))
    print(f"HTML Report generated: {output_file}")

    # Generate PDF
    pdf_file = output_file.replace(".html", ".pdf")
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        font_config = FontConfiguration()
        html_string = template.render(context)
        
        # Create HTML object with proper base URL for loading resources
        html_obj = HTML(string=html_string, base_url=input_dir)
        
        # Define CSS for better print output
        print_css = CSS(string='''
            @page {
                size: A4;
                margin: 20mm 15mm 25mm 15mm;
            }
            @page :first {
                margin: 0;
            }
            body {
                font-family: Georgia, "Times New Roman", serif !important;
            }
        ''', font_config=font_config)
        
        # Generate PDF with print-optimized settings
        html_obj.write_pdf(
            pdf_file,
            stylesheets=[print_css],
            font_config=font_config
        )
        print(f"PDF Report generated: {pdf_file}")
    except Exception as e:
        print(f"Error generating PDF: {e}")

    # Generate DOCX
    docx_file = output_file.replace(".html", ".docx")
    try:
        generate_docx_report(context, docx_file)
        print(f"DOCX Report generated: {docx_file}")
    except Exception as e:
        print(f"Error generating DOCX: {e}")

def generate_docx_report(context, output_path):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # 1. Cover Page
    doc.add_heading('Final Analysis Report', 0)
    p = doc.add_paragraph()
    run = p.add_run(f"Project ID: {context['project_id']}")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(f"Client: {context['client_name']} ({context['client_org']})")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2. Table of Contents (Placeholder)
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph("1. Project Details\n2. Sample Details\n3. Sequencing Stats\n4. Mapping Stats\n5. Differential Expression\n6. Functional Annotation")
    doc.add_page_break()

    # 3. Project Details
    doc.add_heading('1 Project Details', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    meta_rows = [
        ('Project ID', context['project_id']),
        ('Project PI', context['project_pi']),
        ('Application', context['application']),
        ('No. of Samples', context['no_of_samples']),
        ('Reference Organism', context['reference_organism']),
        ('Date', context['report_date'])
    ]
    for field, val in meta_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(field)
        row_cells[1].text = str(val)

    # 4. Sequencing Stats
    doc.add_heading('2 Sequencing Data Statistics', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sample'
    hdr_cells[1].text = 'Total Reads'
    hdr_cells[2].text = 'Total Bases'
    hdr_cells[3].text = 'Total Data (GB)'
    
    for s in context['sequencing_stats']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(s.get('sample', ''))
        row_cells[1].text = str(s.get('total_reads', ''))
        row_cells[2].text = str(s.get('total_bases', ''))
        row_cells[3].text = str(s.get('data_gb', ''))

    # 5. Mapping Stats
    doc.add_heading('3 Mapping Statistics', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sample'
    hdr_cells[1].text = 'Total Reads'
    hdr_cells[2].text = 'Mapped Reads'
    hdr_cells[3].text = '% Mapped'
    hdr_cells[4].text = '% Unique'
    
    for s in context['mapping_stats']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(s.get('sample_name', ''))
        row_cells[1].text = str(s.get('total_reads', ''))
        row_cells[2].text = str(s.get('mapped_reads', ''))
        row_cells[3].text = str(s.get('percent_mapped', ''))
        row_cells[4].text = str(s.get('percent_unique', ''))

    # 6. DGE Stats
    doc.add_heading('4 Differential Expression Statistics', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Comparison'
    hdr_cells[1].text = 'Total DEGs'
    hdr_cells[2].text = 'Upregulated'
    hdr_cells[3].text = 'Downregulated'
    hdr_cells[4].text = 'Total Significant'
    
    for s in context['diff_expr_stats']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(s.get('comparison', ''))
        row_cells[1].text = str(s.get('total_degs', ''))
        row_cells[2].text = str(s.get('sig_up', ''))
        row_cells[3].text = str(s.get('sig_down', ''))
        row_cells[4].text = str(s.get('total_sig', ''))

    # 7. Functional Annotation
    doc.add_heading('5 Functional Annotation', level=1)
    doc.add_heading('5.1 KEGG Pathway Statistics', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Level 1'
    hdr_cells[1].text = 'Level 2'
    hdr_cells[2].text = 'Gene Count'
    
    for p in context['pathway_stats']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(p.get('level1', ''))
        row_cells[1].text = str(p.get('level2', ''))
        row_cells[2].text = str(p.get('count', ''))

    doc.add_heading('5.2 GO Distribution', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ontology'
    hdr_cells[1].text = 'Term'
    hdr_cells[2].text = 'Gene Count'
    
    for g in context['go_distribution']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(g.get('ontology', ''))
        row_cells[1].text = str(g.get('term', ''))
        row_cells[2].text = str(g.get('count', ''))

    doc.save(output_path)

if __name__ == "__main__":
    main()
